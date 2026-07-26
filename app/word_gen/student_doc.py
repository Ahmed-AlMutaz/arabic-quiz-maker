import os
import docx
from app.schemas.exam import GeneratedExam, QuestionType
from app.word_gen.styles import docx_styler, RGBColor
from app.core.logging import logger

class StudentDocGenerator:
    """Generates clean, printable Student Examination Word (.docx) document."""

    @staticmethod
    def generate(exam: GeneratedExam, output_path: str) -> str:
        logger.info("Generating Student docx document...", exam_id=exam.exam_id, output_path=output_path)
        doc = docx.Document()
        docx_styler.apply_document_rtl(doc)

        # Header Box
        docx_styler.create_header_table(doc, exam.metadata)

        # Instructions Callout Box
        p_inst = doc.add_paragraph()
        docx_styler.set_rtl_paragraph(p_inst)
        r_inst_title = p_inst.add_run("تعليمات هامة للطلاب:\n")
        docx_styler.set_rtl_run(r_inst_title, font_size=12, bold=True, color=docx_styler.ACCENT_COLOR)
        r_inst_body = p_inst.add_run("• أجب عن جميع الأسئلة بخط واضح ومقروء.\n• اقرأ الأسئلة بتمعن قبل الإجابة عليها.")
        docx_styler.set_rtl_run(r_inst_body, font_size=11, color=docx_styler.MUTED_TEXT)
        doc.add_paragraph()

        # Questions Rendering Grouped by Type
        questions_by_type = {}
        for q in exam.questions:
            q_type = q.question_type.value if hasattr(q.question_type, "value") else str(q.question_type)
            questions_by_type.setdefault(q_type, []).append(q)

        type_titles = {
            "mcq": "أسئلة الاختيار من متعدد (اختر الإجابة الصحيحة)",
            "true_false": "أجب بصح أو خطأ (ضع علامة ✓ أمام العبارة الصحيحة وعلامة ✗ أمام العبارة الخطأ)",
            "fill_in_blank": "أسئلة أكمل الفراغ (أكمل الفراغات التالية بما يناسبها)",
            "short_answer": "الأسئلة القصيرة (أجب عن الأسئلة التالية)",
            "essay": "الأسئلة المقالية (أسئلة التفكير الناقد والإجابات المفتوحة)"
        }

        q_counter = 1
        for q_type, q_list in questions_by_type.items():
            section_title = type_titles.get(q_type, f"أسئلة {q_type}")
            
            p_sec = doc.add_paragraph()
            docx_styler.set_rtl_paragraph(p_sec)
            r_sec = p_sec.add_run(f"■ {section_title}:")
            docx_styler.set_rtl_run(r_sec, font_size=14, bold=True, color=docx_styler.PRIMARY_COLOR)

            for q in q_list:
                p_q = doc.add_paragraph()
                docx_styler.set_rtl_paragraph(p_q)
                
                # Question text line
                r_num = p_q.add_run(f"({q_counter}) ")
                docx_styler.set_rtl_run(r_num, font_size=13, bold=True, color=docx_styler.DARK_TEXT)
                
                r_qtext = p_q.add_run(f"{q.question_text} ")
                docx_styler.set_rtl_run(r_qtext, font_size=13, bold=False, color=docx_styler.DARK_TEXT)
                
                r_marks = p_q.add_run(f"[{q.marks} درجات]")
                docx_styler.set_rtl_run(r_marks, font_size=11, bold=True, color=docx_styler.MUTED_TEXT)

                # MCQ Options Rendering
                if q_type == "mcq" and q.options:
                    p_opts = doc.add_paragraph()
                    docx_styler.set_rtl_paragraph(p_opts)
                    opt_str = "    ".join([f"({opt.key}) {opt.text}" for opt in q.options])
                    r_opts = p_opts.add_run(f"    {opt_str}")
                    docx_styler.set_rtl_run(r_opts, font_size=12, color=docx_styler.DARK_TEXT)

                # Write-in Answer Lines for Short Answer / Essay
                elif q_type in ["short_answer", "essay"]:
                    for _ in range(2 if q_type == "short_answer" else 4):
                        p_line = doc.add_paragraph()
                        docx_styler.set_rtl_paragraph(p_line)
                        r_line = p_line.add_run("    الإجابة: ............................................................................................................................")
                        docx_styler.set_rtl_run(r_line, font_size=11, color=RGBColor(180, 180, 180))

                q_counter += 1
                doc.add_paragraph()

        # Footer
        p_footer = doc.add_paragraph()
        docx_styler.set_rtl_paragraph(p_footer)
        p_footer.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        r_foot = p_footer.add_run("<<< انتهت الأسئلة - مع أطيب التمنيات بالتوفيق والنجاح >>>")
        docx_styler.set_rtl_run(r_foot, font_size=12, bold=True, color=docx_styler.PRIMARY_COLOR)

        doc.save(output_path)
        logger.info("Saved Student docx exam file", path=output_path)
        return output_path

student_doc_generator = StudentDocGenerator()
