import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

class DocxRTLStyler:
    """OpenXML RTL Helper for Word Document Generation."""

    PRIMARY_COLOR = RGBColor(26, 82, 118)   # Deep Blue (#1A5276)
    ACCENT_COLOR = RGBColor(180, 40, 40)   # Burgundy Red
    DARK_TEXT = RGBColor(33, 33, 33)      # Charcoal Dark
    MUTED_TEXT = RGBColor(100, 100, 100)   # Muted Gray

    @staticmethod
    def set_cell_background(cell, color_hex: str):
        """Sets background shading color for a table cell."""
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    @staticmethod
    def set_rtl_paragraph(paragraph):
        """Forces Right-To-Left (RTL) paragraph layout in OpenXML."""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pPr = paragraph._p.get_or_add_pPr()
        bidi = parse_xml(f'<w:bidi {nsdecls("w")}/>')
        pPr.append(bidi)

    @staticmethod
    def set_rtl_run(run, font_name: str = "Traditional Arabic", font_size: int = 14, bold: bool = False, color: RGBColor = None):
        """Sets RTL font properties for a text run."""
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold = bold
        if color:
            run.font.color.rgb = color

        rPr = run._r.get_or_add_rPr()
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}"/>')
        rPr.append(rFonts)
        rtl = parse_xml(f'<w:rtl {nsdecls("w")}/>')
        rPr.append(rtl)

    @classmethod
    def apply_document_rtl(cls, doc: docx.Document):
        """Configures document page layout for A4 and RTL direction."""
        for section in doc.sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
            
            # Set section RTL
            sectPr = section._sectPr
            bidi = parse_xml(f'<w:bidi {nsdecls("w")}/>')
            sectPr.append(bidi)

    @classmethod
    def create_header_table(cls, doc: docx.Document, metadata) -> docx.table.Table:
        """Creates professional Arabic exam metadata header box."""
        table = doc.add_table(rows=2, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        # Set cell widths
        widths = [Inches(2.2), Inches(2.5), Inches(2.2)]
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = w

        # Center cell (Exam Title & Subject)
        cell_c = table.cell(0, 1)
        p = cell_c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{metadata.subject_name or 'امتحان مادة اللغة العربية'}")
        cls.set_rtl_run(r, font_size=14, bold=True, color=cls.PRIMARY_COLOR)

        # Bottom full-width row for Student Name & Grade box
        cell_bottom = table.cell(1, 0)
        cell_bottom.merge(table.cell(1, 2))
        cls.set_cell_background(cell_bottom, "F2F4F4")
        p = cell_bottom.paragraphs[0]
        cls.set_rtl_paragraph(p)
        r = p.add_run(f"اسم الطالب: ....................................................   الدرجة الكلية: (       / {metadata.total_marks} )")
        cls.set_rtl_run(r, font_size=13, bold=True, color=cls.DARK_TEXT)

        doc.add_paragraph()  # spacing
        return table

docx_styler = DocxRTLStyler()
