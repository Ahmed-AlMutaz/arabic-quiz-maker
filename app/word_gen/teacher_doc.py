import os
import docx
from app.schemas.exam import GeneratedExam
from app.word_gen.styles import docx_styler, RGBColor
from app.core.logging import logger

class TeacherDocGenerator:
    """Generates Teacher Model Answer Key Word (.docx) document."""

    @staticmethod
    def generate(exam: GeneratedExam, output_path: str) -> str:
        logger.info("Generating Teacher Answer Key docx document...", exam_id=exam.exam_id, output_path=output_path)
        doc = docx.Document()
        docx_styler.apply_document_rtl(doc)

        # Header Box
        docx_styler.create_header_table(doc, exam.metadata)

        # Title Badge
        p_badge = doc.add_paragraph()
        docx_styler.set_rtl_paragraph(p_badge)
        p_badge.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        r_badge = p_badge.add_run("--- نموذج الإجابة النموذجي وسلم الدرجات (خاص بالمدرس) ---")
        docx_styler.set_rtl_run(r_badge, font_size=15, bold=True, color=docx_styler.ACCENT_COLOR)
        doc.add_paragraph()

        # Questions & Answer Table
        table = doc.add_table(rows=1, cols=4)
        table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
        
        # Header Row
        hdr_cells = table.rows[0].cells
        headers = ["رقم السؤال", "نص السؤال والخيارات", "الإجابة النموذجية وشرحها", "الدرجات والمراجع"]
        for i, text in enumerate(headers):
            hdr_cells[i].text = text
            docx_styler.set_cell_background(hdr_cells[i], "1A5276")
            p = hdr_cells[i].paragraphs[0]
            docx_styler.set_rtl_paragraph(p)
            r = p.runs[0]
            docx_styler.set_rtl_run(r, font_size=12, bold=True, color=RGBColor(255, 255, 255))

        for idx, q in enumerate(exam.questions, 1):
            row_cells = table.add_row().cells
            
            # Col 0: Q Number & Type
            type_label_map = {
                "mcq": "اختيار من متعدد",
                "true_false": "أجب بصح أو خطأ",
                "fill_in_blank": "أكمل الفراغ",
                "short_answer": "سؤال قصير",
                "essay": "سؤال مقالي"
            }
            q_type_str = q.question_type.value if hasattr(q.question_type, "value") else str(q.question_type)
            type_label = type_label_map.get(q_type_str, q_type_str)

            p0 = row_cells[0].paragraphs[0]
            docx_styler.set_rtl_paragraph(p0)
            r0 = p0.add_run(f"س {idx}\n({type_label})")
            docx_styler.set_rtl_run(r0, font_size=11, bold=True, color=docx_styler.PRIMARY_COLOR)

            # Col 1: Question Text
            p1 = row_cells[1].paragraphs[0]
            docx_styler.set_rtl_paragraph(p1)
            r1 = p1.add_run(q.question_text)
            docx_styler.set_rtl_run(r1, font_size=11, color=docx_styler.DARK_TEXT)

            # Col 2: Correct Answer & Explanation
            p2 = row_cells[2].paragraphs[0]
            docx_styler.set_rtl_paragraph(p2)
            r2_ans = p2.add_run(f"الإجابة: {q.correct_answer}\n")
            docx_styler.set_rtl_run(r2_ans, font_size=11, bold=True, color=docx_styler.ACCENT_COLOR)
            r2_exp = p2.add_run(f"التوضيح: {q.explanation}")
            docx_styler.set_rtl_run(r2_exp, font_size=10, color=docx_styler.MUTED_TEXT)

            # Col 3: Marks & Context Ref
            p3 = row_cells[3].paragraphs[0]
            docx_styler.set_rtl_paragraph(p3)
            r3 = p3.add_run(f"الدرجة: {q.marks}\nمرجع: {q.context_chunk_id or 'نص الدرس'}")
            docx_styler.set_rtl_run(r3, font_size=10, color=docx_styler.DARK_TEXT)

            # Zebra striping
            if idx % 2 == 0:
                for c in row_cells:
                    docx_styler.set_cell_background(c, "F8F9F9")

        doc.save(output_path)
        logger.info("Saved Teacher docx answer key file", path=output_path)
        return output_path

teacher_doc_generator = TeacherDocGenerator()
